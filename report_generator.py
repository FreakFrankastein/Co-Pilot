"""
Report Generator
==================
Membaca findings.json (hasil dari Burp extension + orchestrator Nuclei/sqlmap)
dan menyusun laporan pentest profesional dalam format .docx.

Cara pakai:
    python report_generator.py --output "Laporan_Pentest_ClientX.docx"

Opsi:
    --client "Nama Client"     -> nama client di cover page
    --scope "Deskripsi scope"  -> deskripsi ruang lingkup pengujian
    --tester "Nama Kamu"       -> nama pentester di cover page
    --include-false-positive   -> sertakan juga finding berstatus false_positive (default: tidak)
"""

import argparse
import json
import os
from collections import defaultdict
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FINDINGS_PATH = "findings.json"

SEVERITY_COLOR = {
    "Critical": RGBColor(0x8B, 0x00, 0x00),
    "critical": RGBColor(0x8B, 0x00, 0x00),
    "High": RGBColor(0xC0, 0x39, 0x2B),
    "high": RGBColor(0xC0, 0x39, 0x2B),
    "Medium": RGBColor(0xD6, 0x8A, 0x00),
    "medium": RGBColor(0xD6, 0x8A, 0x00),
    "Low": RGBColor(0x2E, 0x86, 0xC1),
    "low": RGBColor(0x2E, 0x86, 0xC1),
    "None": RGBColor(0x5D, 0x6D, 0x7E),
    "info": RGBColor(0x5D, 0x6D, 0x7E),
}

SEVERITY_ORDER = ["Critical", "critical", "High", "high", "Medium", "medium",
                  "Low", "low", "None", "info"]


def _shade_cell(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _severity_rank(sev):
    order = ["critical", "high", "medium", "low", "info", "none"]
    sev_l = (sev or "info").lower()
    return order.index(sev_l) if sev_l in order else len(order)


def load_findings(include_false_positive=False):
    if not os.path.exists(FINDINGS_PATH):
        return []
    with open(FINDINGS_PATH, "r") as f:
        findings = json.load(f)
    if not include_false_positive:
        findings = [f for f in findings if f.get("status") != "false_positive"]
    findings.sort(key=lambda f: _severity_rank(f.get("severity") or f.get("cvss_severity")))
    return findings


def add_cover_page(doc, client, scope, tester):
    doc.add_paragraph().add_run().add_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LAPORAN PENETRATION TESTING")
    run.bold = True
    run.font.size = Pt(28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(client or "[Nama Client]")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph().add_run().add_break()
    doc.add_paragraph().add_run().add_break()

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ("Tanggal Laporan", datetime.now().strftime("%d %B %Y")),
        ("Pentester", tester or "[Nama Pentester]"),
        ("Ruang Lingkup", scope or "[Deskripsi scope pengujian]"),
        ("Klasifikasi", "CONFIDENTIAL"),
    ]
    for i, (label, value) in enumerate(rows_data):
        meta_table.cell(i, 0).text = label
        meta_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        meta_table.cell(i, 1).text = str(value)

    doc.add_page_break()


def add_executive_summary(doc, findings):
    doc.add_heading("1. Executive Summary", level=1)

    counts = defaultdict(int)
    for f in findings:
        sev = (f.get("severity") or f.get("cvss_severity") or "info")
        counts[sev.lower() if isinstance(sev, str) else "info"] += 1

    p = doc.add_paragraph()
    p.add_run(
        f"Pengujian penetrasi telah dilakukan dan berhasil mengidentifikasi "
        f"total {len(findings)} temuan yang telah diverifikasi. Ringkasan "
        f"distribusi tingkat keparahan (severity) sebagai berikut:"
    )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Severity"
    hdr[1].text = "Jumlah Temuan"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    for sev_label in ["critical", "high", "medium", "low", "info"]:
        if counts.get(sev_label, 0) > 0:
            row = table.add_row().cells
            row[0].text = sev_label.capitalize()
            row[1].text = str(counts[sev_label])
            color = SEVERITY_COLOR.get(sev_label)
            if color:
                row[0].paragraphs[0].runs[0].font.color.rgb = color
                row[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()


def add_methodology(doc):
    doc.add_heading("2. Metodologi", level=1)
    doc.add_paragraph(
        "Pengujian dilakukan menggunakan kombinasi automated scanning dan "
        "manual testing, mencakup:"
    )
    items = [
        "Traffic interception & passive analysis via Burp Suite",
        "Automated vulnerability scanning (template-based) menggunakan Nuclei",
        "Active SQL injection testing menggunakan sqlmap",
        "Manual verification terhadap seluruh temuan otomatis sebelum dikonfirmasi",
        "Manual business-logic testing (autentikasi, otorisasi, alur transaksi)",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Skema penilaian risiko menggunakan CVSS versi 4.0 sesuai standar "
        "FIRST.org."
    )
    doc.add_paragraph()


def add_findings_detail(doc, findings):
    doc.add_heading("3. Detail Temuan", level=1)

    if not findings:
        doc.add_paragraph("Tidak ada temuan yang tercatat.")
        return

    for idx, f in enumerate(findings, start=1):
        sev = (f.get("severity") or f.get("cvss_severity") or "info")
        sev_display = sev.capitalize() if isinstance(sev, str) else "Info"

        heading = doc.add_heading(level=2)
        run = heading.add_run(f"3.{idx} {f.get('name') or f.get('subtype') or 'Unnamed Finding'}")

        # Severity badge line
        badge_p = doc.add_paragraph()
        badge_run = badge_p.add_run(f"Severity: {sev_display}")
        badge_run.bold = True
        badge_run.font.color.rgb = SEVERITY_COLOR.get(sev, RGBColor(0, 0, 0))

        table = doc.add_table(rows=0, cols=2)
        table.style = "Light List Accent 1"

        def add_row(label, value):
            row = table.add_row().cells
            row[0].text = label
            row[0].paragraphs[0].runs[0].bold = True
            row[1].text = str(value) if value else "-"

        add_row("Sumber Daya Terdampak (Endpoint)", f.get("endpoint"))
        add_row("Parameter Rentan", f.get("parameter"))
        add_row("Metode HTTP", f.get("method"))
        add_row("Tipe Injeksi", f.get("injection_type"))
        add_row("Sumber Deteksi", f.get("source", "manual/heuristic"))
        add_row("Status Verifikasi", f.get("status", "candidate"))
        add_row("CVSS 4.0 Vector", f.get("cvss_vector") or f.get("suggested_cvss_vector"))
        add_row("CVSS Score", f.get("cvss_score"))
        add_row("Evidence", f.get("evidence"))
        add_row("Catatan", f.get("note"))

        # PoC section
        poc = f.get("poc")
        if poc:
            doc.add_paragraph()
            poc_heading = doc.add_paragraph()
            poc_run = poc_heading.add_run("Proof of Concept (PoC)")
            poc_run.bold = True
            poc_run.font.size = Pt(11)

            # PoC command (pakai monospace font)
            if poc.get("command"):
                cmd_label = doc.add_paragraph()
                cmd_label.add_run("Command:").bold = True
                cmd_para = doc.add_paragraph()
                cmd_run = cmd_para.add_run(poc["command"])
                cmd_run.font.name = "Courier New"
                cmd_run.font.size = Pt(9)

            # Tool yang dipakai
            if poc.get("tool"):
                tool_p = doc.add_paragraph()
                tool_p.add_run("Tool: ").bold = True
                tool_p.add_run(poc["tool"])

            # Langkah validasi manual
            if poc.get("steps"):
                steps_label = doc.add_paragraph()
                steps_label.add_run("Langkah Validasi Manual:").bold = True
                for step in poc["steps"]:
                    doc.add_paragraph(step, style="List Bullet")

            # WAF note (kalau ada)
            if poc.get("waf_note"):
                waf_p = doc.add_paragraph()
                waf_run = waf_p.add_run(f"⚠️  Catatan WAF: {poc['waf_note']}")
                waf_run.font.color.rgb = RGBColor(0xD6, 0x8A, 0x00)

        doc.add_paragraph()


def add_recommendations(doc):
    doc.add_heading("4. Rekomendasi Umum", level=1)
    recs = [
        "Terapkan parameterized query / prepared statement untuk seluruh akses database.",
        "Lakukan validasi & encoding output secara konsisten untuk mencegah XSS.",
        "Hilangkan verbose error message dan stack trace dari response ke client.",
        "Terapkan secret management yang layak (vault) untuk API key & credential.",
        "Lakukan re-test setelah remediasi untuk memastikan temuan telah tertutup.",
    ]
    for r in recs:
        doc.add_paragraph(r, style="List Bullet")


def generate_report(output_path, client, scope, tester, include_false_positive):
    findings = load_findings(include_false_positive)

    doc = Document()
    add_cover_page(doc, client, scope, tester)
    add_executive_summary(doc, findings)
    add_methodology(doc)
    add_findings_detail(doc, findings)
    add_recommendations(doc)

    doc.save(output_path)
    print(f"[+] Laporan berhasil dibuat: {output_path}")
    print(f"[+] Total temuan disertakan: {len(findings)}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate laporan pentest dari findings.json")
    parser.add_argument("--output", default="Laporan_Pentest.docx")
    parser.add_argument("--client", default="")
    parser.add_argument("--scope", default="")
    parser.add_argument("--tester", default="")
    parser.add_argument("--include-false-positive", action="store_true")
    args = parser.parse_args()

    generate_report(args.output, args.client, args.scope, args.tester,
                     args.include_false_positive)
